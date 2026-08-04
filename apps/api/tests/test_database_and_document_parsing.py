import base64
import json
import sqlite3
from datetime import UTC, datetime
from pathlib import Path

import pytest

from app.core.database import metadata
from app.modules.auth.domain import User
from app.modules.auth.infra.sqlite_users import SQLiteUserRepository
from app.modules.materials.application.material_parser import parse_material_to_candidates
from app.modules.materials.domain import UploadedMaterialRecord
from app.modules.materials.infra.sqlite_store import MaterialSQLiteStore


def test_production_schema_contains_auth_material_version_and_graph_tables() -> None:
    table_names = set(metadata.tables)

    assert {
        "users",
        "roles",
        "permissions",
        "user_roles",
        "role_permissions",
        "uploaded_materials",
        "material_versions",
        "academic_support_links",
        "recognition_candidates",
        "graph_nodes",
        "graph_edges",
        "diagnostic_findings",
        "improvement_tasks",
    } <= table_names


@pytest.mark.asyncio
async def test_sqlite_user_repository_persists_users_roles_and_permissions(
    tmp_path: Path,
) -> None:
    repo = SQLiteUserRepository(base_dir=tmp_path)

    admin = await repo.get_by_username("admin")
    assert admin is not None
    assert admin.role == "admin"

    created = await repo.create(
        User(
            id="user-material-owner",
            username="material-owner",
            password_hash="hashed",
            display_name="Material Owner",
            role="teacher",
            created_at=datetime.now(UTC).isoformat(),
        )
    )
    assert created.username == "material-owner"
    assert await repo.get_by_id("user-material-owner") is not None

    with sqlite3.connect(tmp_path / "ea_mvp.sqlite3") as conn:
        role_count = conn.execute("select count(*) from roles").fetchone()[0]
        permission_count = conn.execute("select count(*) from permissions").fetchone()[0]
        user_role = conn.execute(
            """
            select role_id from user_roles
            where user_id = 'user-material-owner'
            """
        ).fetchone()

    assert role_count >= 2
    assert permission_count >= 5
    assert user_role[0] == "role-teacher"


def test_material_store_creates_and_updates_material_version(tmp_path: Path) -> None:
    store = MaterialSQLiteStore(base_dir=tmp_path)
    record = store.create_from_base64(
        user_id="user-1",
        uploaded_by="Teacher",
        file_name="course-outline.txt",
        category="course_outline",
        content_base64=base64.b64encode("课程目标 C-05-01".encode()).decode(),
        content_type="text/plain",
        course="Embedded Systems",
    )

    with sqlite3.connect(tmp_path / "ea_mvp.sqlite3") as conn:
        version = conn.execute(
            """
            select version_no, checksum, parsed_artifact_json
            from material_versions
            where material_id = ?
            """,
            (record.id,),
        ).fetchone()
    assert version[0] == 1
    assert len(version[1]) == 64
    assert version[2] == "{}"

    artifact = {"parser": {"strategy": "rules+catalog"}}
    store.mark_parsed(
        user_id="user-1",
        material_id=record.id,
        extracted_text="课程目标 C-05-01",
        extracted_node_count=1,
        candidates_created=1,
        parser_version="hybrid-material-parser v0.7",
        parse_strategy="rules+catalog",
        parsed_artifact_json=json.dumps(artifact, ensure_ascii=False),
    )

    with sqlite3.connect(tmp_path / "ea_mvp.sqlite3") as conn:
        updated = conn.execute(
            """
            select parser_version, parse_strategy, parsed_artifact_json
            from material_versions
            where material_id = ?
            """,
            (record.id,),
        ).fetchone()

    assert updated[0] == "hybrid-material-parser v0.7"
    assert updated[1] == "rules+catalog"
    assert json.loads(updated[2]) == artifact


def test_docx_table_structure_is_extracted(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "outline.docx"
    document = Document()
    document.add_paragraph("课程名称：嵌入式系统原理")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "实验项目"
    table.cell(0, 1).text = "支撑指标点"
    table.cell(1, 0).text = "GPIO 与定时器综合实验"
    table.cell(1, 1).text = "C-05-01"
    document.save(path)

    result = parse_material_to_candidates(_record_for(path, "docx"))

    extraction = result.structured_artifact["extraction"]
    assert extraction["strategy"] == "python-docx"
    assert extraction["tableCount"] == 1
    assert "GPIO 与定时器综合实验" in result.text
    assert result.structured_artifact["capability"]["tableExtraction"] is True


def test_xlsx_sheet_rows_and_merged_ranges_are_extracted(tmp_path: Path) -> None:
    from openpyxl import Workbook

    path = tmp_path / "rubric.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "评分项"
    sheet["A1"] = "实验项目"
    sheet["B1"] = "指标点"
    sheet["C1"] = "分值"
    sheet["A2"] = "传感器数据采集实验"
    sheet["B2"] = "C-05-01"
    sheet["C2"] = 20
    sheet.merge_cells("A4:C4")
    sheet["A4"] = "课程目标与毕业要求支撑说明"
    workbook.save(path)
    workbook.close()

    result = parse_material_to_candidates(_record_for(path, "xlsx"))

    extraction = result.structured_artifact["extraction"]
    assert extraction["strategy"] == "openpyxl"
    assert extraction["sheetCount"] == 1
    assert extraction["tableCount"] == 1
    assert extraction["sheets"][0]["mergedRanges"] == ["A4:C4"]
    assert "传感器数据采集实验" in result.text


def _record_for(path: Path, file_type: str) -> UploadedMaterialRecord:
    now = datetime.now(UTC).isoformat()
    return UploadedMaterialRecord(
        id=f"material-{file_type}",
        user_id="user-test",
        file_name=path.name,
        category="test",
        content_type="application/octet-stream",
        file_type=file_type,
        size_bytes=path.stat().st_size,
        stored_path=str(path),
        status="pending",
        uploaded_by="Tester",
        created_at=now,
        updated_at=now,
        course="嵌入式系统原理",
    )
