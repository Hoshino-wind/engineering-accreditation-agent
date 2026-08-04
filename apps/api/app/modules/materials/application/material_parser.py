import re
import zipfile
from dataclasses import dataclass
from datetime import datetime
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from app.modules.academic.domain import (
    AcademicCatalog,
    CompetencyIndicator,
    Course,
    ExperimentProject,
)
from app.modules.llm.domain.models import ExtractionItem, RelationItem
from app.modules.materials.domain import UploadedMaterialRecord
from app.modules.recognition.domain.candidate import (
    CandidateEvidence,
    RecognitionCandidate,
    RecognitionCandidateRisk,
    RecognitionCandidateType,
)


@dataclass(frozen=True, slots=True)
class ParsedNode:
    id: str
    kind: str
    code: str
    name: str
    description: str
    confidence: float
    source_excerpt: str


@dataclass(frozen=True, slots=True)
class ParseResult:
    text: str
    nodes: list[ParsedNode]
    candidates: list[RecognitionCandidate]
    structured_artifact: dict[str, Any]
    extraction_artifact: dict[str, Any]
    parser_version: str = "hybrid-material-parser v0.7"
    parse_strategy: str = "rules+catalog"


@dataclass(frozen=True, slots=True)
class IndicatorMatch:
    indicator: CompetencyIndicator
    confidence: int
    risk: RecognitionCandidateRisk
    reason: str


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    text: str
    artifact: dict[str, Any]


def get_ocr_runtime_status() -> dict[str, Any]:
    try:
        import pytesseract
    except ImportError as exc:
        return {
            "available": False,
            "status": "python_dependency_missing",
            "engine": "tesseract",
            "version": None,
            "languages": [],
            "message": f"pytesseract is not installed: {exc}",
        }
    try:
        version = str(pytesseract.get_tesseract_version())
        languages = sorted(pytesseract.get_languages(config=""))
        return {
            "available": True,
            "status": "ready",
            "engine": "tesseract",
            "version": version,
            "languages": languages,
            "message": "OCR runtime is available.",
        }
    except Exception as exc:  # noqa: BLE001
        return {
            "available": False,
            "status": "engine_unavailable",
            "engine": "tesseract",
            "version": None,
            "languages": [],
            "message": str(exc),
        }


def parse_material_to_candidates(
    record: UploadedMaterialRecord,
    catalog: AcademicCatalog | None = None,
) -> ParseResult:
    document = _extract_document(record)
    text = document.text
    if not text.strip():
        text = f"{record.file_name}\n{record.category}\n{record.course or ''}"

    course = _match_course(record, text, catalog)
    nodes = _extract_nodes(record, text, catalog, course)
    candidates = _build_candidates(record, text, nodes, catalog, course)
    structured = _build_structured_artifact(
        record=record,
        text=text,
        nodes=nodes,
        candidates=candidates,
        course=course,
        parse_strategy="rules+catalog",
        llm=None,
        extraction=document.artifact,
    )
    return ParseResult(
        text=text[:20000],
        nodes=nodes,
        candidates=candidates,
        structured_artifact=structured,
        extraction_artifact=document.artifact,
    )


def merge_llm_parse_result(
    *,
    record: UploadedMaterialRecord,
    base: ParseResult,
    catalog: AcademicCatalog | None,
    extracted_items: list[ExtractionItem],
    relation_items: list[RelationItem],
    llm_model: str,
) -> ParseResult:
    llm_nodes = _nodes_from_llm(record, extracted_items)
    nodes = _dedupe_nodes([*base.nodes, *llm_nodes])
    candidates = _dedupe_candidates(
        [
            *base.candidates,
            *_candidates_from_llm_relations(
                record=record,
                text=base.text,
                nodes=nodes,
                relation_items=relation_items,
                catalog=catalog,
                llm_model=llm_model,
            ),
        ]
    )
    structured = _build_structured_artifact(
        record=record,
        text=base.text,
        nodes=nodes,
        candidates=candidates,
        course=None,
        parse_strategy="rules+catalog+llm",
        llm={
            "model": llm_model,
            "extractedNodeCount": len(extracted_items),
            "relationCount": len(relation_items),
        },
        extraction=base.extraction_artifact,
    )
    return ParseResult(
        text=base.text,
        nodes=nodes,
        candidates=candidates,
        structured_artifact=structured,
        extraction_artifact=base.extraction_artifact,
        parser_version="hybrid-material-parser v0.7",
        parse_strategy="rules+catalog+llm",
    )


def _nodes_from_llm(
    record: UploadedMaterialRecord,
    items: list[ExtractionItem],
) -> list[ParsedNode]:
    nodes: list[ParsedNode] = []
    for index, item in enumerate(items, start=1):
        nodes.append(
            ParsedNode(
                id=f"{record.id}-llm-{index}",
                kind=_llm_kind_to_node_kind(item.kind),
                code=item.code or f"LLM-{index:02d}",
                name=item.name,
                description=item.description or "",
                confidence=item.confidence,
                source_excerpt=item.source_excerpt or "",
            )
        )
    return nodes


def _candidates_from_llm_relations(
    *,
    record: UploadedMaterialRecord,
    text: str,
    nodes: list[ParsedNode],
    relation_items: list[RelationItem],
    catalog: AcademicCatalog | None,
    llm_model: str,
) -> list[RecognitionCandidate]:
    if not catalog:
        return []
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    excerpt = _excerpt(text)
    candidates: list[RecognitionCandidate] = []
    for index, relation in enumerate(relation_items, start=1):
        indicator = _indicator_from_relation(relation, catalog)
        if indicator is None:
            continue
        source = _node_from_relation(relation, nodes)
        confidence = _relation_confidence(relation.confidence)
        candidates.append(
            RecognitionCandidate(
                id=f"candidate-{record.id}-llm-{indicator.code.lower()}-{index}",
                title=f"{record.file_name} LLM 支撑关系候选 {index}",
                course=record.course or _course_name_from_catalog(catalog),
                candidate_type=RecognitionCandidateType.RELATION,
                confidence=confidence,
                risk=_risk_from_confidence(confidence),
                source_node=source.name if source else relation.source_id,
                relation=relation.relation_type,
                target_node=f"能力指标 {indicator.code} {indicator.title}",
                explanation=(
                    f"LLM inferred this relation using {llm_model}. "
                    f"Reasoning: {relation.reasoning or 'No reasoning returned.'}"
                ),
                processor_version="hybrid-material-parser v0.7",
                generated_at=generated_at,
                impact_course_objectives=_objective_count_for_course(catalog, None),
                impact_ability_nodes=1,
                impact_rubric_items=max(_rubric_count_for_indicator(catalog, indicator.id), 1),
                support_strength=relation.strength,
                evidence=(
                    CandidateEvidence(
                        id=f"evidence-{record.id}-llm-{indicator.code.lower()}-{index}",
                        resource_name=record.file_name,
                        resource_version="uploaded+llm",
                        coordinate=f"LLM:{relation.source_id}->{relation.target_id}",
                        excerpt=excerpt,
                        hash=f"SHA256 {sha256(excerpt.encode('utf-8')).hexdigest()[:12]}",
                    ),
                ),
            )
        )
    return candidates


def _build_structured_artifact(
    *,
    record: UploadedMaterialRecord,
    text: str,
    nodes: list[ParsedNode],
    candidates: list[RecognitionCandidate],
    course: Course | None,
    parse_strategy: str,
    llm: dict[str, Any] | None,
    extraction: dict[str, Any],
) -> dict[str, Any]:
    suffix = Path(record.file_name).suffix.lower()
    table_count = int(extraction.get("tableCount") or 0)
    ocr = extraction.get("ocr") if isinstance(extraction.get("ocr"), dict) else {}
    warnings = extraction.get("warnings") if isinstance(extraction.get("warnings"), list) else []
    return {
        "parser": {
            "version": "hybrid-material-parser v0.7",
            "strategy": parse_strategy,
            "sources": ["rules", "academic_catalog", *(["llm"] if llm else [])],
            "llm": llm,
        },
        "material": {
            "id": record.id,
            "fileName": record.file_name,
            "fileType": record.file_type,
            "category": record.category,
            "course": course.name if course else record.course,
            "textChars": len(text),
        },
        "capability": {
            "textExtraction": bool(text.strip()),
            "tableExtraction": table_count > 0,
            "ocrRequired": bool(ocr.get("required")),
            "ocrStatus": ocr.get("status", "not_applicable"),
            "warnings": warnings,
            "notes": _parser_capability_notes(suffix, extraction),
        },
        "extraction": extraction,
        "nodes": [
            {
                "id": node.id,
                "kind": node.kind,
                "code": node.code,
                "name": node.name,
                "confidence": node.confidence,
                "sourceExcerpt": node.source_excerpt,
            }
            for node in nodes
        ],
        "candidateRelations": [
            {
                "id": candidate.id,
                "sourceNode": candidate.source_node,
                "relation": candidate.relation,
                "targetNode": candidate.target_node,
                "confidence": candidate.confidence,
                "strength": candidate.support_strength
                or _strength_from_confidence(candidate.confidence),
                "evidenceHash": candidate.evidence[0].hash if candidate.evidence else None,
            }
            for candidate in candidates
        ],
        "detectedIndicatorCodes": _unique(
            [
                code
                for candidate in candidates
                for code in re.findall(r"C-\d{2}-\d{2}", candidate.target_node)
            ]
        ),
    }


def _llm_kind_to_node_kind(kind: str) -> str:
    return {
        "course": "Course",
        "experiment": "Experiment",
        "knowledge": "KnowledgePoint",
        "resource": "TeachingResource",
    }.get(kind, "TeachingResource")


def _indicator_from_relation(
    relation: RelationItem,
    catalog: AcademicCatalog,
) -> CompetencyIndicator | None:
    target = relation.target_id.casefold()
    for indicator in catalog.indicators:
        if target in {indicator.id.casefold(), indicator.code.casefold()}:
            return indicator
    return None


def _node_from_relation(
    relation: RelationItem,
    nodes: list[ParsedNode],
) -> ParsedNode | None:
    source = relation.source_id.casefold()
    for node in nodes:
        if source in {node.id.casefold(), node.code.casefold(), node.name.casefold()}:
            return node
    return None


def _relation_confidence(value: float) -> int:
    confidence = value * 100 if value <= 1 else value
    return max(1, min(99, round(confidence)))


def _risk_from_confidence(confidence: int) -> RecognitionCandidateRisk:
    if confidence >= 86:
        return RecognitionCandidateRisk.HIGH_IMPACT
    if confidence < 76:
        return RecognitionCandidateRisk.LOW_CONFIDENCE
    return RecognitionCandidateRisk.NORMAL


def _strength_from_confidence(confidence: int) -> str:
    if confidence >= 85:
        return "strong"
    if confidence >= 70:
        return "medium"
    return "weak"


def _course_name_from_catalog(catalog: AcademicCatalog) -> str:
    return catalog.courses[0].name if catalog.courses else "Unknown course"


def _dedupe_nodes(nodes: list[ParsedNode]) -> list[ParsedNode]:
    seen: set[tuple[str, str]] = set()
    result: list[ParsedNode] = []
    for node in nodes:
        key = (node.kind, node.code.casefold() or node.name.casefold())
        if key in seen:
            continue
        seen.add(key)
        result.append(node)
    return result


def _dedupe_candidates(
    candidates: list[RecognitionCandidate],
) -> list[RecognitionCandidate]:
    seen: set[tuple[str, str]] = set()
    result: list[RecognitionCandidate] = []
    for candidate in candidates:
        key = (candidate.source_node.casefold(), candidate.target_node.casefold())
        if key in seen:
            continue
        seen.add(key)
        result.append(candidate)
    return result


def _parser_capability_notes(suffix: str, extraction: dict[str, Any]) -> list[str]:
    if suffix == ".pdf":
        ocr = extraction.get("ocr") if isinstance(extraction.get("ocr"), dict) else {}
        notes = ["PDF text extraction uses pypdf and keeps page-level trace metadata."]
        if ocr.get("required"):
            notes.append("The file appears to be scanned or image-heavy; OCR was attempted.")
        if ocr.get("status") in {"unavailable", "failed"}:
            notes.append("Install Tesseract OCR and language packs before production OCR use.")
        return notes
    if suffix == ".xlsx":
        return [
            "Spreadsheet sheets, cell coordinates, row previews and merged ranges are extracted.",
            "Formula values use workbook cached values when available.",
        ]
    if suffix == ".docx":
        return [
            "Word paragraphs and tables are extracted with row and cell previews.",
            "Images inside Word files are marked for future OCR if needed.",
        ]
    return ["Plain text extraction is supported."]


def _extract_document(record: UploadedMaterialRecord) -> ExtractedDocument:
    path = Path(record.stored_path)
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text, encoding = _extract_plain_text(path)
        return ExtractedDocument(
            text=text,
            artifact={
                "format": suffix.lstrip(".") or "txt",
                "strategy": "plain-text",
                "encoding": encoding,
                "textChars": len(text),
                "tableCount": 0,
                "warnings": [],
            },
        )
    if suffix == ".docx":
        return _extract_docx_document(path)
    if suffix == ".xlsx":
        return _extract_xlsx_document(path)
    if suffix == ".pdf":
        return _extract_pdf_document(path, record.file_name)
    return ExtractedDocument(
        text=f"{record.file_name}\n{record.category}",
        artifact={
            "format": suffix.lstrip(".") or "unknown",
            "strategy": "file-metadata-fallback",
            "textChars": len(record.file_name) + len(record.category),
            "tableCount": 0,
            "warnings": ["Unsupported file type; parser used file metadata only."],
        },
    )


def _extract_plain_text(path: Path) -> tuple[str, str]:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            return raw.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore"), "utf-8-ignore"


def _extract_docx_document(path: Path) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError:
        text = _extract_docx_text_fallback(path)
        return ExtractedDocument(
            text=text,
            artifact={
                "format": "docx",
                "strategy": "ooxml-fallback",
                "textChars": len(text),
                "tableCount": 0,
                "warnings": ["python-docx is not installed; used raw OOXML text fallback."],
            },
        )

    document = Document(str(path))
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    tables: list[dict[str, Any]] = []
    table_lines: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            values = [_clean_cell(cell.text) for cell in row.cells]
            if any(values):
                rows.append(values)
                table_lines.append(f"DOCX Table {table_index}: " + " | ".join(values))
        column_count = max((len(row) for row in rows), default=0)
        tables.append(
            {
                "index": table_index,
                "rows": len(rows),
                "columns": column_count,
                "preview": rows[:8],
            }
        )
    text = "\n".join([*paragraphs, *table_lines])
    return ExtractedDocument(
        text=text,
        artifact={
            "format": "docx",
            "strategy": "python-docx",
            "paragraphCount": len(paragraphs),
            "tableCount": len(tables),
            "tables": tables,
            "textChars": len(text),
            "warnings": [],
        },
    )


def _extract_docx_text_fallback(path: Path) -> str:
    with zipfile.ZipFile(path) as docx:
        xml = docx.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    texts = [
        elem.text
        for elem in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        if elem.text
    ]
    return "\n".join(texts)


def _extract_xlsx_document(path: Path) -> ExtractedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError:
        text = _extract_xlsx_text_fallback(path)
        return ExtractedDocument(
            text=text,
            artifact={
                "format": "xlsx",
                "strategy": "ooxml-fallback",
                "textChars": len(text),
                "tableCount": 0,
                "warnings": ["openpyxl is not installed; used raw OOXML cell fallback."],
            },
        )

    workbook = load_workbook(path, data_only=True, read_only=False)
    sheets: list[dict[str, Any]] = []
    text_lines: list[str] = []
    table_count = 0
    for worksheet in workbook.worksheets:
        rows: list[list[str]] = []
        non_empty_cells = 0
        for row in worksheet.iter_rows():
            values = [_cell_to_text(cell.value) for cell in row]
            if not any(values):
                continue
            non_empty_cells += sum(1 for value in values if value)
            compact = _trim_empty_tail(values)
            rows.append(compact)
            text_lines.append(f"{worksheet.title}: " + " | ".join(compact))
        merged_ranges = [str(item) for item in worksheet.merged_cells.ranges]
        if rows:
            table_count += 1
        sheets.append(
            {
                "name": worksheet.title,
                "maxRow": worksheet.max_row,
                "maxColumn": worksheet.max_column,
                "nonEmptyRows": len(rows),
                "nonEmptyCells": non_empty_cells,
                "mergedRanges": merged_ranges,
                "preview": rows[:10],
            }
        )
    workbook.close()
    text = "\n".join(text_lines[:2000])
    return ExtractedDocument(
        text=text,
        artifact={
            "format": "xlsx",
            "strategy": "openpyxl",
            "sheetCount": len(sheets),
            "tableCount": table_count,
            "sheets": sheets,
            "textChars": len(text),
            "warnings": [],
        },
    )


def _extract_xlsx_text_fallback(path: Path) -> str:
    texts: list[str] = []
    with zipfile.ZipFile(path) as workbook:
        shared_strings = _read_xlsx_shared_strings(workbook)
        for name in workbook.namelist():
            if name == "xl/sharedStrings.xml":
                continue
            if not name.startswith("xl/worksheets/"):
                continue
            try:
                root = ElementTree.fromstring(workbook.read(name))
            except ElementTree.ParseError:
                continue
            for cell in root.iter():
                tag = _strip_namespace(cell.tag)
                if tag != "c":
                    continue
                cell_type = cell.attrib.get("t")
                value = next(
                    (
                        child.text
                        for child in cell
                        if _strip_namespace(child.tag) == "v"
                    ),
                    None,
                )
                if not value:
                    continue
                if cell_type == "s" and value.isdigit():
                    index = int(value)
                    if 0 <= index < len(shared_strings):
                        texts.append(shared_strings[index])
                else:
                    texts.append(value)
    return "\n".join(texts[:1000])


def _extract_pdf_document(path: Path, file_name: str) -> ExtractedDocument:
    warnings: list[str] = []
    pages: list[dict[str, Any]] = []
    text_lines: list[str] = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        for index, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:  # noqa: BLE001
                page_text = ""
                warnings.append(f"Page {index} text extraction failed: {exc}")
            cleaned = page_text.strip()
            pages.append({"page": index, "textChars": len(cleaned)})
            if cleaned:
                text_lines.append(f"PDF Page {index}\n{cleaned}")
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"pypdf extraction failed: {exc}")

    text = "\n".join(text_lines)
    ocr_required = len(text.strip()) < 120
    ocr_result: dict[str, Any] = {
        "required": ocr_required,
        "status": "not_required" if not ocr_required else "pending",
        "pagesAttempted": 0,
        "textChars": 0,
    }
    if ocr_required:
        ocr_text, ocr_result = _try_pdf_ocr(path)
        if ocr_text.strip():
            text = "\n".join([text, ocr_text]).strip()
        elif not text.strip():
            text = (
                f"{file_name}\n"
                "PDF appears to be scanned or image-based. OCR did not produce text."
            )
    return ExtractedDocument(
        text=text,
        artifact={
            "format": "pdf",
            "strategy": "pypdf+optional-ocr",
            "pageCount": len(pages),
            "pages": pages,
            "tableCount": 0,
            "ocr": ocr_result,
            "textChars": len(text),
            "warnings": warnings,
        },
    )


def _try_pdf_ocr(path: Path, max_pages: int = 3) -> tuple[str, dict[str, Any]]:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        return "", {
            "required": True,
            "status": "unavailable",
            "pagesAttempted": 0,
            "textChars": 0,
            "reason": f"OCR dependency is not installed: {exc}",
        }

    texts: list[str] = []
    pages_attempted = 0
    try:
        with fitz.open(path) as document:
            for page_index in range(min(max_pages, document.page_count)):
                page = document.load_page(page_index)
                pages_attempted += 1
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.open(BytesIO(pixmap.tobytes("png")))
                page_text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
                if page_text:
                    texts.append(f"OCR Page {page_index + 1}\n{page_text}")
    except Exception as exc:  # noqa: BLE001
        return "", {
            "required": True,
            "status": "failed",
            "pagesAttempted": pages_attempted,
            "textChars": 0,
            "reason": str(exc),
        }

    text = "\n".join(texts)
    return text, {
        "required": True,
        "status": "success" if text.strip() else "empty",
        "pagesAttempted": pages_attempted,
        "textChars": len(text),
    }


def _read_xlsx_shared_strings(workbook: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in workbook.namelist():
        return []
    try:
        root = ElementTree.fromstring(workbook.read("xl/sharedStrings.xml"))
    except ElementTree.ParseError:
        return []
    values: list[str] = []
    for item in root.iter():
        if _strip_namespace(item.tag) != "si":
            continue
        values.append(
            "".join(
                text.text or ""
                for text in item.iter()
                if _strip_namespace(text.tag) == "t"
            )
        )
    return values


def _clean_cell(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    return _clean_cell(str(value))


def _trim_empty_tail(values: list[str]) -> list[str]:
    trimmed = list(values)
    while trimmed and not trimmed[-1]:
        trimmed.pop()
    return trimmed


def _strip_namespace(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _extract_nodes(
    record: UploadedMaterialRecord,
    text: str,
    catalog: AcademicCatalog | None,
    course: Course | None,
) -> list[ParsedNode]:
    excerpt = _excerpt(text)
    course_name = course.name if course else record.course or _guess_course_name(record, text)
    nodes = [
        ParsedNode(
            id=f"{record.id}-course",
            kind="Course",
            code=course.code if course else _code_from_text(course_name, "COURSE"),
            name=course_name,
            description=f"Course/material node extracted from {record.file_name}.",
            confidence=0.9 if course else 0.78,
            source_excerpt=excerpt,
        )
    ]

    matched_experiments = _match_experiments(text, catalog, course)
    for index, experiment in enumerate(matched_experiments[:5], start=1):
        nodes.append(
            ParsedNode(
                id=f"{record.id}-exp-{index}",
                kind="Experiment",
                code=experiment.code,
                name=experiment.title,
                description=experiment.description,
                confidence=0.9,
                source_excerpt=excerpt,
            )
        )

    if len(nodes) == 1:
        for index, name in enumerate(_find_experiment_names(text)[:3], start=1):
            nodes.append(
                ParsedNode(
                    id=f"{record.id}-exp-{index}",
                    kind="Experiment",
                    code=f"EXP-{index:02d}",
                    name=name,
                    description="Experiment/task node extracted from uploaded material.",
                    confidence=0.76,
                    source_excerpt=excerpt,
                )
            )

    return nodes


def _build_candidates(
    record: UploadedMaterialRecord,
    text: str,
    nodes: list[ParsedNode],
    catalog: AcademicCatalog | None,
    course: Course | None,
) -> list[RecognitionCandidate]:
    source_node = next((n for n in nodes if n.kind == "Experiment"), nodes[0])
    matches = _infer_indicator_matches(text, catalog)
    if not matches:
        matches = _fallback_indicator_matches(text)

    excerpt = _excerpt(text)
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    candidates: list[RecognitionCandidate] = []

    for index, match in enumerate(matches[:5], start=1):
        indicator = match.indicator
        rubric_count = _rubric_count_for_indicator(catalog, indicator.id)
        target_node = f"能力指标 {indicator.code} {indicator.title}".strip()
        candidates.append(
            RecognitionCandidate(
                id=f"candidate-{record.id}-{indicator.code.lower()}-{index}",
                title=f"{record.file_name} 支撑关系候选 {index}",
                course=course.name if course else record.course or nodes[0].name,
                candidate_type=RecognitionCandidateType.RELATION,
                confidence=match.confidence,
                risk=match.risk,
                source_node=source_node.name,
                relation="supports",
                target_node=target_node,
                explanation=(
                    f"{match.reason} Source material: {record.file_name}. "
                    "The relation must be reviewed by a teacher before it is "
                    "projected into the official graph."
                ),
                processor_version="mvp-file-parser v0.6",
                generated_at=generated_at,
                impact_course_objectives=_objective_count_for_course(catalog, course),
                impact_ability_nodes=1,
                impact_rubric_items=max(rubric_count, 1),
                evidence=(
                    CandidateEvidence(
                        id=f"evidence-{record.id}-{indicator.code.lower()}-{index}",
                        resource_name=record.file_name,
                        resource_version="uploaded",
                        coordinate=_evidence_coordinate(source_node, indicator),
                        excerpt=excerpt,
                        hash=f"SHA256 {sha256(excerpt.encode('utf-8')).hexdigest()[:12]}",
                    ),
                ),
            )
        )
    return candidates


def _match_course(
    record: UploadedMaterialRecord,
    text: str,
    catalog: AcademicCatalog | None,
) -> Course | None:
    if not catalog:
        return None
    haystack = _normalize_text(
        " ".join([record.file_name, record.course or "", text[:4000]])
    )
    best: tuple[int, Course] | None = None
    for course in catalog.courses:
        score = 0
        if course.code and _normalize_text(course.code) in haystack:
            score += 5
        if course.name and _normalize_text(course.name) in haystack:
            score += 8
        score += _keyword_overlap(
            haystack,
            _normalize_text(f"{course.name} {course.category} {course.owner}"),
        )
        if score and (best is None or score > best[0]):
            best = (score, course)
    if best:
        return best[1]
    return catalog.courses[0] if len(catalog.courses) == 1 else None


def _match_experiments(
    text: str,
    catalog: AcademicCatalog | None,
    course: Course | None,
) -> list[ExperimentProject]:
    if not catalog:
        return []
    haystack = _normalize_text(text)
    rows = [
        experiment
        for experiment in catalog.experiments
        if course is None or experiment.course_id == course.id
    ]
    scored: list[tuple[int, ExperimentProject]] = []
    for experiment in rows:
        score = 0
        if _normalize_text(experiment.code) in haystack:
            score += 5
        if experiment.title and _normalize_text(experiment.title) in haystack:
            score += 8
        score += _keyword_overlap(
            haystack,
            _normalize_text(
                f"{experiment.title} {experiment.description} {experiment.environment}"
            ),
        )
        if score:
            scored.append((score, experiment))
    return [experiment for _, experiment in sorted(scored, key=lambda item: item[0], reverse=True)]


def _infer_indicator_matches(
    text: str,
    catalog: AcademicCatalog | None,
) -> list[IndicatorMatch]:
    if not catalog:
        return []
    haystack = _normalize_text(text)
    explicit_codes = {
        _normalize_code(code)
        for code in re.findall(r"C-\d{2}-\d{2}", text, flags=re.I)
    }
    scored: list[tuple[int, IndicatorMatch]] = []

    for indicator in catalog.indicators:
        score = 0
        normalized_code = _normalize_code(indicator.code)
        if normalized_code in explicit_codes:
            score += 12
        if _normalize_text(indicator.code) in haystack:
            score += 8
        descriptor = _normalize_text(f"{indicator.title} {indicator.description}")
        score += _keyword_overlap(haystack, descriptor)
        if score <= 0:
            continue
        confidence = min(94, 62 + score * 3)
        risk = RecognitionCandidateRisk.NORMAL
        if normalized_code in explicit_codes or confidence >= 86:
            risk = RecognitionCandidateRisk.HIGH_IMPACT
        elif confidence < 76:
            risk = RecognitionCandidateRisk.LOW_CONFIDENCE
        reason = (
            f"Matched catalog indicator {indicator.code} by explicit code."
            if normalized_code in explicit_codes
            else f"Matched catalog indicator {indicator.code} by keyword overlap."
        )
        scored.append((score, IndicatorMatch(indicator, confidence, risk, reason)))

    return [match for _, match in sorted(scored, key=lambda item: item[0], reverse=True)]


def _fallback_indicator_matches(text: str) -> list[IndicatorMatch]:
    targets = _infer_targets(text)
    return [
        IndicatorMatch(
            indicator=CompetencyIndicator(
                id=f"fallback-{_normalize_code(code).lower()}",
                requirement_id="fallback",
                code=code,
                title=title,
                description=title,
            ),
            confidence=confidence,
            risk=risk,
            reason="Matched by fallback parser rules.",
        )
        for code, title, confidence, risk in targets
    ]


def _infer_targets(text: str) -> list[tuple[str, str, int, RecognitionCandidateRisk]]:
    explicit = _unique(re.findall(r"C-\d{2}-\d{2}", text, flags=re.I))
    if explicit:
        return [
            (
                _normalize_code(code),
                "课程材料显式引用的指标点",
                88,
                RecognitionCandidateRisk.HIGH_IMPACT,
            )
            for code in explicit[:3]
        ]

    lowered = text.lower()
    inferred: list[tuple[str, str, int, RecognitionCandidateRisk]] = []
    if any(word in lowered for word in ("embedded", "stm32", "gpio", "fpga", "verilog", "adc")):
        inferred.append(("C-05-01", "现代工具选择与使用", 82, RecognitionCandidateRisk.NORMAL))
        inferred.append(("C-03-01", "系统设计方法", 74, RecognitionCandidateRisk.LOW_CONFIDENCE))
    if any(word in lowered for word in ("数据结构", "链表", "排序", "算法", "complexity")):
        inferred.append(("C-01-01", "工程知识应用", 86, RecognitionCandidateRisk.NORMAL))
        inferred.append(("C-01-02", "问题推演与分析", 76, RecognitionCandidateRisk.LOW_CONFIDENCE))
    if not inferred:
        inferred.append(("C-01-01", "工程知识应用", 72, RecognitionCandidateRisk.LOW_CONFIDENCE))
    return inferred


def _rubric_count_for_indicator(catalog: AcademicCatalog | None, indicator_id: str) -> int:
    if not catalog:
        return 0
    return sum(1 for item in catalog.rubric_items if item.indicator_id == indicator_id)


def _objective_count_for_course(catalog: AcademicCatalog | None, course: Course | None) -> int:
    if not catalog or not course:
        return 1
    count = sum(1 for item in catalog.objectives if item.course_id == course.id)
    return max(count, 1)


def _evidence_coordinate(source_node: ParsedNode, indicator: CompetencyIndicator) -> str:
    return f"{source_node.kind}:{source_node.code} -> Indicator:{indicator.code}"


def _guess_course_name(record: UploadedMaterialRecord, text: str) -> str:
    found = _find_first(
        text,
        [
            r"课程名称[:：]\s*([^\n\r]{2,40})",
            r"Course\s*Name[:：]\s*([^\n\r]{2,60})",
        ],
    )
    return found or Path(record.file_name).stem[:40]


def _find_experiment_names(text: str) -> list[str]:
    values = _find_all(
        text,
        [
            r"实验[一二三四五六七八九十\d]+[:：\s]*([^\n\r]{2,50})",
            r"(GPIO[^\n\r]{0,30}|STM32[^\n\r]{0,30}|ADC[^\n\r]{0,30}|链表[^\n\r]{0,30}|排序[^\n\r]{0,30})",
        ],
    )
    return [value.strip(" ：:，,。.;；") for value in values]


def _find_first(text: str, patterns: list[str]) -> str | None:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return None


def _find_all(text: str, patterns: list[str]) -> list[str]:
    values: list[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            value = match.group(1) if match.lastindex else match.group(0)
            value = value.strip()
            if value and value not in values:
                values.append(value)
    return values


def _excerpt(text: str) -> str:
    collapsed = re.sub(r"\s+", " ", text).strip()
    return collapsed[:300] or "Material uploaded. Waiting for teacher review."


def _code_from_text(text: str, prefix: str) -> str:
    digest = sha256(text.encode("utf-8")).hexdigest()[:6].upper()
    return f"{prefix}-{digest}"


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text).casefold()


def _normalize_code(code: str) -> str:
    return code.upper()


def _keyword_overlap(haystack: str, descriptor: str) -> int:
    if not descriptor:
        return 0
    score = 0
    for token in _tokens(descriptor):
        if len(token) < 2:
            continue
        if token in haystack:
            score += 1
    return min(score, 8)


def _tokens(text: str) -> list[str]:
    english = re.findall(r"[a-zA-Z][a-zA-Z0-9_+-]{1,}", text.casefold())
    chinese = re.findall(r"[\u4e00-\u9fff]{2,}", text)
    chunks: list[str] = []
    for value in chinese:
        chunks.extend(value[index : index + 4] for index in range(0, max(len(value) - 1, 0), 2))
    return _unique([*english, *chunks])


def _unique(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        normalized = value.upper()
        if normalized in seen:
            continue
        seen.add(normalized)
        result.append(value)
    return result
